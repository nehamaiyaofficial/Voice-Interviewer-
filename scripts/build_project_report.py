from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "VoiceHire_NLP_Project_Report.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(33, 37, 41)
MUTED = RGBColor(95, 99, 104)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_cell_text(cell, bold=False, color=INK, size=10.5) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_FILL)
        style_cell_text(cell, bold=True, color=INK, size=10)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            style_cell_text(cells[idx], size=10)
    set_table_width(table, widths)
    document.add_paragraph()
    return table


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(11)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.size = Pt(10.5)
    set_table_width(table, [6.5])
    document.add_paragraph()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("VoiceHire-NLP Project Report")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("VoiceHire-NLP")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BLUE

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Adaptive Voice Interview Simulator with Multi-Axis Evaluation")
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Professional Project Report")
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    document.add_paragraph()
    add_callout(
        document,
        "Project Summary",
        "VoiceHire-NLP is a Streamlit-based mock interview system that speaks interview questions, records student voice answers, transcribes them with Whisper, evaluates each answer across multiple dimensions, asks adaptive follow-up questions, and generates a final feedback report after a 10-question session.",
    )

    metadata = [
        ["Project Type", "NLP course project / mock interview simulator"],
        ["Core Stack", "Python, Streamlit, faster-whisper, spaCy, scikit-learn"],
        ["Primary Output", "Interview log, final feedback report, demo recording"],
        ["Evaluation Axes", "Relevance, structure, clarity, confidence, keyword coverage, fluency"],
    ]
    add_table(document, ["Field", "Details"], metadata, [1.8, 4.7])
    document.add_page_break()


def build_report() -> None:
    document = Document()
    configure_document(document)
    add_title_page(document)

    document.add_heading("1. Abstract", level=1)
    document.add_paragraph(
        "VoiceHire-NLP is an adaptive voice interview simulator designed to demonstrate practical NLP concepts through a complete interview workflow. The system generates interview questions, speaks them using browser text-to-speech, records spoken answers, transcribes them using Whisper, evaluates each answer across multiple dimensions, and asks follow-up questions that visibly reference the student's previous response. After ten questions, the system generates a structured feedback report containing overall performance, NLP analysis, communication analysis, suggested topics, and failure analysis."
    )

    document.add_heading("2. Project Objectives", level=1)
    objectives = [
        "Build a complete 10-question mock interview workflow with persistent state.",
        "Use voice input and speech-to-text transcription instead of typed answers.",
        "Evaluate answers on relevance, structure, clarity, confidence, keyword coverage, and fluency.",
        "Demonstrate course NLP concepts including Bag of Words, TF-IDF, N-Gram analysis, POS tagging, NER, classification-style scoring, and evaluation metrics.",
        "Generate adaptive follow-up questions that explicitly refer to the previous answer.",
        "Produce a final report and support a complete recorded demo session.",
    ]
    for item in objectives:
        add_bullet(document, item)

    document.add_heading("3. System Architecture", level=1)
    document.add_paragraph(
        "The application is organized as a modular Python project. Streamlit provides the user interface, the interviewer module manages question generation, the speech-to-text module handles Whisper transcription, the evaluator module performs multi-axis scoring, and the report generator creates the final feedback report."
    )
    architecture_rows = [
        ["1", "Role input", "Student selects the target job role."],
        ["2", "Question generation", "The system creates the first interview question and later adaptive follow-ups."],
        ["3", "TTS playback", "The Speak question button uses browser speech synthesis to read the question aloud."],
        ["4", "Voice answer", "The student records an answer using the Streamlit microphone input."],
        ["5", "Whisper transcription", "faster-whisper converts the recorded audio into text."],
        ["6", "NLP analysis", "BoW, TF-IDF, N-Gram, POS, NER, and audio features are extracted."],
        ["7", "Evaluation", "Each answer is scored on six axes and stored in interview memory."],
        ["8", "Final report", "After ten questions, the app generates a feedback report and JSON log."],
    ]
    add_table(document, ["Step", "Component", "Purpose"], architecture_rows, [0.55, 1.65, 4.3])

    document.add_heading("4. Functional Features", level=1)
    feature_rows = [
        ["Voice question", "Browser TTS speaks each generated question."],
        ["Voice answer", "Student answers through microphone recording only."],
        ["Whisper transcription", "Local faster-whisper model transcribes speech to text."],
        ["Adaptive follow-up", "Next question references a keyword or idea from the previous answer."],
        ["State management", "InterviewMemory stores all ten turns, evaluations, and audio metadata."],
        ["Final report", "Report includes scores, strengths, weaknesses, NLP analysis, communication analysis, suggested topics, and failure analysis."],
    ]
    add_table(document, ["Feature", "Implementation"], feature_rows, [1.8, 4.7])

    document.add_heading("5. NLP Lab Concept Mapping", level=1)
    lab_rows = [
        ["Bag of Words", "Keyword coverage score by matching expected role/question keywords."],
        ["TF-IDF", "Important concept extraction from current and previous answer transcripts."],
        ["N-Gram analysis", "Fluency score based on repetition and bigram/trigram diversity."],
        ["POS tagging", "Sentence structure analysis using nouns, verbs, and adjectives."],
        ["NER", "Detection of technologies, tools, companies, and skill entities."],
        ["Text classification", "Good answer vs needs improvement label based on scoring threshold."],
        ["Evaluation metrics", "Aggregated answer scores and final percentage report."],
    ]
    add_table(document, ["Lab Concept", "Use in Project"], lab_rows, [1.8, 4.7])

    document.add_heading("6. Evaluation Methodology", level=1)
    document.add_paragraph(
        "Each answer is evaluated out of 60 marks. The scoring model combines transcript-based NLP features, communication indicators, and rule-based scoring logic. This approach makes the evaluation transparent and easy to explain during a project demonstration."
    )
    metric_rows = [
        ["Relevance", "10", "Checks whether the answer covers expected role and question keywords."],
        ["Structure", "10", "Rewards organized answers and balanced sentence structure."],
        ["Clarity", "10", "Considers answer length, punctuation, and repetition."],
        ["Confidence", "10", "Uses speaking pace, filler words, and answer length as approximate indicators."],
        ["Keyword Coverage", "10", "Measures Bag of Words keyword match percentage."],
        ["Fluency", "10", "Uses N-Gram diversity and repetition analysis."],
    ]
    add_table(document, ["Metric", "Marks", "Evaluation Basis"], metric_rows, [1.5, 0.8, 4.2])

    document.add_heading("7. Audio and Communication Analysis", level=1)
    document.add_paragraph(
        "The system performs a separate audio communication analysis pass. It estimates speaking pace from transcript length and recording duration, counts filler words such as um, uh, like, actually, and basically, and tracks pause indicators where available. These features support the bonus communication analysis requirement."
    )
    for item in [
        "Speaking pace is reported as words per minute.",
        "Filler word count is included in each answer evaluation and final report.",
        "Low-quality recordings are detected using duration and volume checks before transcription.",
        "The final report summarizes average answer length, total filler words, and average speaking pace.",
    ]:
        add_bullet(document, item)

    document.add_heading("8. Adaptive Questioning", level=1)
    document.add_paragraph(
        "Adaptive questioning is implemented in the interviewer module. After each answer, the system selects a meaningful word or suggested topic from the transcript and includes it in the next question. This makes the follow-up visibly connected to the student's previous response."
    )
    add_callout(
        document,
        "Example",
        "If the student mentions inheritance, the next question may begin with: 'You mentioned inheritance. Explain one important technical concept you would use as a Python Developer. Please connect it with polymorphism.'",
    )

    document.add_heading("9. Report Generation and Submission Artifacts", level=1)
    document.add_paragraph(
        "The project produces the artifacts required for submission. The Streamlit app can save a JSON interview log and a final text report after the 10-question session. The notebook demonstrates multi-turn state and report generation. The demo video must be recorded manually during a complete run of the application."
    )
    artifact_rows = [
        ["Notebook", "notebooks/VoiceHire_Final.ipynb", "Documents the pipeline and demonstrates multi-turn state."],
        ["Interview log", "logs/interview_log.json", "Generated after a completed app session."],
        ["Final report", "reports/final_report.txt", "Generated feedback report after ten questions."],
        ["Demo recording", "session_recording.mp4", "Manual screen/audio recording of one complete session."],
        ["Recorded answers", "audio/a1.wav ... audio/a10.wav", "Saved voice answers for the interview turns."],
    ]
    add_table(document, ["Artifact", "Path", "Purpose"], artifact_rows, [1.5, 2.1, 2.9])

    document.add_heading("10. Failure Analysis", level=1)
    document.add_paragraph(
        "Failure analysis is important because automated interview evaluation contains subjective and technical uncertainty. The system includes explicit failure-analysis notes in the final report."
    )
    failure_rows = [
        ["Subjective scoring", "A keyword-heavy weak answer may receive a high score.", "Use human review alongside automated feedback."],
        ["Transcription error", "Whisper may mishear technical terms such as inheritance as appearance.", "Show transcript and allow re-recording if speech quality is poor."],
        ["Confidence scoring", "Pace and filler words do not fully represent true confidence.", "Treat confidence as an approximate communication signal."],
        ["Audio quality", "Quiet or short recordings may not transcribe correctly.", "Check duration and volume before transcription."],
    ]
    add_table(document, ["Issue", "Example Risk", "Mitigation"], failure_rows, [1.45, 2.65, 2.4])

    document.add_heading("11. Testing and Verification", level=1)
    document.add_paragraph(
        "The project includes automated tests for the evaluator and audio-quality helper. The application has also been checked through Python compilation and Streamlit startup verification. The README includes Git Bash and PowerShell instructions so the project works both locally and on GitHub."
    )
    for item in [
        "pytest test suite validates evaluator output and audio-quality handling.",
        "compileall verifies Python syntax across app, source modules, and tests.",
        "GitHub Actions workflow runs tests on push and pull request.",
        "Generated model files, recordings, virtual environments, and logs are ignored by Git unless intentionally submitted.",
    ]:
        add_bullet(document, item)

    document.add_heading("12. Conclusion", level=1)
    document.add_paragraph(
        "VoiceHire-NLP satisfies the required mock interview workflow and demonstrates NLP course concepts in a practical, end-to-end system. It combines voice interaction, speech transcription, classical NLP feature extraction, adaptive state management, multi-axis evaluation, and final report generation. The project is suitable for demonstration as a complete NLP application rather than a generic chatbot because each report section explicitly connects implementation behavior to the underlying lab concepts."
    )

    document.add_heading("Appendix A: File Structure", level=1)
    file_rows = [
        ["app.py", "Main Streamlit application and user workflow."],
        ["src/interviewer.py", "Initial and adaptive follow-up question generation."],
        ["src/speech_to_text.py", "Recording quality check and Whisper transcription."],
        ["src/evaluator.py", "Multi-axis answer scoring."],
        ["src/bow_analysis.py", "Keyword coverage and tokenization."],
        ["src/tfidf_analysis.py", "Important term extraction."],
        ["src/ngram_analysis.py", "Fluency and repetition analysis."],
        ["src/pos_ner_analysis.py", "POS and NER analysis."],
        ["src/audio_analysis.py", "Speaking pace and filler word analysis."],
        ["src/report_generator.py", "Final feedback report generation."],
    ]
    add_table(document, ["File", "Role"], file_rows, [2.1, 4.4])

    document.add_heading("Appendix B: Demonstration Checklist", level=1)
    checklist = [
        "Start the Streamlit app.",
        "Click Speak question and confirm the question is spoken.",
        "Record each voice answer for at least 8 seconds.",
        "Submit the answer and confirm the evaluation appears.",
        "Confirm the next question references the previous answer.",
        "Complete all 10 questions.",
        "Save interview_log.json and final_report.txt.",
        "Record the complete session as session_recording.mp4.",
    ]
    for item in checklist:
        add_number(document, item)

    OUT.parent.mkdir(exist_ok=True)
    document.save(OUT)
    print("reports/VoiceHire_NLP_Project_Report.docx")


if __name__ == "__main__":
    build_report()
